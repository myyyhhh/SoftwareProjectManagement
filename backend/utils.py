import base64
import io
import os
from typing import Optional

import docx
import fitz  # PyMuPDF
import requests
from lxml import html
from openai import OpenAI



SUPPORTED_FILE_EXTENSIONS = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".webp"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}


def _clean_text(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip()).strip()


def _extract_docx_tables(doc: docx.Document) -> str:
    table_blocks = []
    for idx, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            table_blocks.append(f"[DOCX表格 {idx}]\n" + "\n".join(rows))
    return "\n\n".join(table_blocks)


def _extract_pdf_tables(doc: fitz.Document) -> str:
    table_blocks = []
    for page_index, page in enumerate(doc, start=1):
        try:
            finder = page.find_tables()
            tables = getattr(finder, "tables", []) if finder else []
        except Exception:
            tables = []

        for table_index, table in enumerate(tables, start=1):
            try:
                data = table.extract()
            except Exception:
                data = []

            rows = []
            for row in data or []:
                clean_row = [str(cell).strip().replace("\n", " ") if cell is not None else "" for cell in row]
                if any(clean_row):
                    rows.append(" | ".join(clean_row))
            if rows:
                table_blocks.append(f"[PDF第{page_index}页表格 {table_index}]\n" + "\n".join(rows))

    return "\n\n".join(table_blocks)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    tables = _extract_docx_tables(doc)
    joined = "\n".join(paragraphs)
    if tables:
        joined += f"\n\n{tables}"
    return _clean_text(joined)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = []
    for page_index, page in enumerate(doc, start=1):
        page_text = page.get_text("text") or ""
        if page_text.strip():
            text_parts.append(f"[PDF第{page_index}页]\n{page_text.strip()}")

    tables = _extract_pdf_tables(doc)
    if tables:
        text_parts.append(tables)

    return _clean_text("\n\n".join(text_parts))


def extract_text_from_image_with_vlm(file_bytes: bytes, filename: str) -> str:
    """
    用多模态模型做轻量 OCR / 图片理解。
    默认优先使用千问兼容接口（QWEN_* 环境变量），也兼容 OPENAI_* 变量兜底。
    """
    api_key = os.getenv("QWEN_API_KEY") or os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError(
            "图片内容提取需要配置 QWEN_API_KEY（或兼容的 OPENAI_API_KEY）"
        )

    model = (
        os.getenv("QWEN_VLM_MODEL")
        or os.getenv("MULTIMODAL_MODEL")
        or os.getenv("OPENAI_MODEL")
        or "qwen-vl-max-latest"
    )
    base_url = (
        os.getenv("QWEN_BASE_URL")
        or os.getenv("OPENAI_BASE_URL")
        or "https://dashscope.aliyuncs.com/compatible-mode/v1"
    )
    client = OpenAI(api_key=api_key, base_url=base_url)

    ext = os.path.splitext(filename.lower())[1]
    mime = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
    }.get(ext, "image/png")

    encoded = base64.b64encode(file_bytes).decode("utf-8")
    prompt = (
        "你是一个严谨的信息提取助手。请尽量完整提取图像中的文字、表格关键信息、标题、参数、卖点和限制条件。"
        "如果无法确认某部分，请明确标注为‘图像中无法可靠识别’，不要臆造。"
    )

    response = client.chat.completions.create(
        model=model,
        temperature=0,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime};base64,{encoded}"},
                    },
                ],
            }
        ],
    )

    content = response.choices[0].message.content or ""
    return _clean_text(content)


def fetch_webpage_text(url: str, timeout: int = 15) -> str:
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/133.0.0.0 Safari/537.36"
        )
    }
    response = requests.get(url, headers=headers, timeout=timeout)
    response.raise_for_status()

    doc = html.fromstring(response.content)
    for bad in doc.xpath("//script|//style|//noscript"):
        bad.drop_tree()

    chunks = []
    title = " ".join(doc.xpath("//title/text()"))
    if title.strip():
        chunks.append(f"[网页标题]\n{title.strip()}")

    meta_desc = " ".join(doc.xpath("//meta[@name='description']/@content"))
    if meta_desc.strip():
        chunks.append(f"[网页描述]\n{meta_desc.strip()}")

    headings = [t.strip() for t in doc.xpath("//h1//text() | //h2//text() | //h3//text()") if t.strip()]
    if headings:
        chunks.append("[网页标题层级]\n" + "\n".join(headings[:50]))

    paragraphs = [t.strip() for t in doc.xpath("//p//text()") if t.strip()]
    if paragraphs:
        chunks.append("[网页正文]\n" + "\n".join(paragraphs[:200]))

    table_blocks = []
    for idx, table in enumerate(doc.xpath("//table"), start=1):
        rows = []
        for row in table.xpath(".//tr"):
            cells = [" ".join(cell.xpath(".//text() ")).strip() for cell in row.xpath("./th|./td")]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            table_blocks.append(f"[网页表格 {idx}]\n" + "\n".join(rows))
    if table_blocks:
        chunks.append("\n\n".join(table_blocks))

    text = _clean_text("\n\n".join(chunks))
    if not text:
        raise ValueError("无法从网页中提取到有效文本")
    return text


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if ext == ".docx":
        return extract_text_from_docx(file_bytes)
    if ext in IMAGE_EXTENSIONS:
        return extract_text_from_image_with_vlm(file_bytes, filename)
    raise ValueError("不支持的文件格式，仅支持 .pdf / .docx / .png / .jpg / .jpeg / .webp")


def build_source_material(
    file_text: Optional[str] = None,
    source_url: Optional[str] = None,
    raw_text: Optional[str] = None,
) -> str:
    blocks = []

    if raw_text and raw_text.strip():
        blocks.append(f"[用户直接输入的产品资料]\n{raw_text.strip()}")

    if source_url and source_url.strip():
        url_text = fetch_webpage_text(source_url.strip())
        blocks.append(f"[产品网页来源]\nURL: {source_url.strip()}\n{url_text}")

    if file_text and file_text.strip():
        blocks.append(f"[上传文件来源]\n{file_text.strip()}")

    combined = _clean_text("\n\n".join(blocks))
    if not combined:
        raise ValueError("至少需要提供一种有效输入：文件、网页 URL 或原始文本")
    return combined
